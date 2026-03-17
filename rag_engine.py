"""
RAG Engine — All settings come from config.py. Don't edit this unless you're changing logic.
"""

import os, re, math, base64
from pathlib import Path
from collections import Counter

import config as cfg
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_chroma import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_core.messages import HumanMessage
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
import fitz  # PyMuPDF
from docx import Document as DocxDoc


# ── BM25 (lightweight, no extra dependency) ──────────────────────────

class SimpleBM25:
    """Minimal BM25 implementation. No external library needed."""

    def __init__(self, k1: float = 1.5, b: float = 0.75):
        self.k1, self.b = k1, b
        self.docs, self.doc_freqs, self.avg_dl, self.corpus_size = [], Counter(), 0, 0

    @staticmethod
    def _tokenize(text: str) -> list[str]:
        return re.findall(r"\w+", text.lower())

    def fit(self, documents: list[str]):
        self.docs = [self._tokenize(d) for d in documents]
        self.corpus_size = len(self.docs)
        self.avg_dl = sum(len(d) for d in self.docs) / max(self.corpus_size, 1)
        self.doc_freqs = Counter()
        for doc in self.docs:
            for term in set(doc):
                self.doc_freqs[term] += 1

    def score(self, query: str) -> list[float]:
        q_terms = self._tokenize(query)
        scores = []
        for doc in self.docs:
            s = 0.0
            doc_len = len(doc)
            tf_map = Counter(doc)
            for t in q_terms:
                if t not in self.doc_freqs:
                    continue
                tf = tf_map.get(t, 0)
                df = self.doc_freqs[t]
                idf = math.log((self.corpus_size - df + 0.5) / (df + 0.5) + 1)
                num = tf * (self.k1 + 1)
                den = tf + self.k1 * (1 - self.b + self.b * doc_len / max(self.avg_dl, 1))
                s += idf * num / den
            scores.append(s)
        return scores


# ── Main RAG Class ───────────────────────────────────────────────────

class MultimodalRAG:
    def __init__(self, api_key: str):
        os.environ["OPENAI_API_KEY"] = api_key

        self.llm = ChatOpenAI(
            model=cfg.LLM_MODEL, temperature=cfg.LLM_TEMPERATURE,
            max_tokens=cfg.LLM_MAX_TOKENS
        )
        self.embeddings = OpenAIEmbeddings(model=cfg.EMBEDDING_MODEL)
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=cfg.CHUNK_SIZE, chunk_overlap=cfg.CHUNK_OVERLAP,
            separators=cfg.SEPARATORS
        )

        chroma = lambda name: Chroma(
            collection_name=name, embedding_function=self.embeddings,
            persist_directory=cfg.CHROMA_PATH,
            collection_metadata={"hnsw:space": "cosine"}
        )
        self.vectorstore = chroma(cfg.COLLECTION_DOCS)
        self.cache_store = chroma(cfg.COLLECTION_CACHE)
        self.bm25 = SimpleBM25()
        self.all_chunks: list[str] = []
        self.all_metas: list[dict] = []
        self._rebuild_bm25()
        self.bts = {}

    def _rebuild_bm25(self):
        """Sync BM25 index with ChromaDB contents."""
        try:
            data = self.vectorstore._collection.get()
            if data and data["documents"]:
                self.all_chunks = data["documents"]
                self.all_metas = data["metadatas"] if data["metadatas"] else [{}] * len(self.all_chunks)
                self.bm25.fit(self.all_chunks)
        except Exception:
            self.all_chunks, self.all_metas = [], []

    # ── File Processing ──────────────────────────────────────────────

    def _describe_image(self, b64: str, ext: str) -> str:
        mime = f"image/{'jpeg' if ext in ('jpg', 'jpeg') else ext}"
        msg = HumanMessage(content=[
            {"type": "text", "text": cfg.IMAGE_DESCRIPTION_PROMPT},
            {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}}
        ])
        return self.llm.invoke([msg]).content

    def _extract_pdf(self, data: bytes) -> list[str]:
        doc = fitz.open(stream=data, filetype="pdf")
        texts, full_text = [], ""
        for page in doc:
            full_text += page.get_text()
            for img in page.get_images(full=True):
                try:
                    base_img = doc.extract_image(img[0])
                    b64 = base64.b64encode(base_img["image"]).decode()
                    texts.append(self._describe_image(b64, base_img.get("ext", "png")))
                except Exception:
                    continue
        doc.close()
        if full_text.strip():
            texts.insert(0, full_text)
        return texts

    def _extract_docx(self, data: bytes) -> str:
        import io
        return "\n".join(p.text for p in DocxDoc(io.BytesIO(data)).paragraphs if p.text.strip())

    def _extract_image(self, data: bytes, ext: str) -> str:
        b64 = base64.b64encode(data).decode()
        return self._describe_image(b64, ext)

    # ── Ingest (from bytes — no file paths) ──────────────────────────

    def ingest(self, filename: str, file_bytes: bytes) -> dict:
        ext = filename.rsplit(".", 1)[-1].lower()

        if ext == "pdf":
            raw_texts = self._extract_pdf(file_bytes)
        elif ext in ("docx", "doc"):
            raw_texts = [self._extract_docx(file_bytes)]
        elif ext in ("png", "jpg", "jpeg", "gif", "webp"):
            raw_texts = [self._extract_image(file_bytes, ext)]
        elif ext == "txt":
            raw_texts = [file_bytes.decode("utf-8", errors="ignore")]
        else:
            return {"error": f"Unsupported file type: .{ext}"}

        chunks = []
        for t in raw_texts:
            chunks.extend(self.splitter.split_text(t))

        docs = [
            Document(page_content=c, metadata={"source": filename, "chunk_id": i})
            for i, c in enumerate(chunks)
        ]
        self.vectorstore.add_documents(docs)
        self._rebuild_bm25()

        return {"file": filename, "chunks": len(chunks), "sources": len(raw_texts)}

    # ── Hybrid Search + Query ────────────────────────────────────────

    def query(self, question: str) -> dict:
        self.bts = {
            "query": question, "top_k": cfg.TOP_K,
            "chunk_size": cfg.CHUNK_SIZE, "chunk_overlap": cfg.CHUNK_OVERLAP,
            "embedding_model": cfg.EMBEDDING_MODEL, "embedding_dim": cfg.EMBEDDING_DIMENSIONS,
            "llm_model": cfg.LLM_MODEL,
            "bm25_weight": cfg.BM25_WEIGHT, "semantic_weight": cfg.SEMANTIC_WEIGHT,
        }

        # ── Step 1: Semantic cache ───────────────────────────────────
        if cfg.CACHE_ENABLED:
            try:
                hits = self.cache_store.similarity_search_with_score(question, k=1)
                if hits and hits[0][1] < (1 - cfg.CACHE_THRESHOLD):
                    self.bts.update(cache_hit=True, cache_similarity=round(1 - hits[0][1], 4))
                    meta = hits[0][0].metadata
                    return {"answer": meta.get("answer", ""), "sources": meta.get("sources", ""), "bts": self.bts}
            except Exception:
                pass
        self.bts["cache_hit"] = False

        # ── Step 2: Semantic search (vector) ─────────────────────────
        total = len(self.all_chunks)
        self.bts["total_chunks_in_db"] = total

        sem_results = self.vectorstore.similarity_search_with_score(question, k=min(cfg.TOP_K * 3, max(total, 1)))
        sem_map = {}
        for doc, dist in sem_results:
            key = doc.page_content[:100]
            if key not in sem_map:
                sem_map[key] = {"doc": doc, "score": 1 - dist}

        # ── Step 3: BM25 keyword search ──────────────────────────────
        bm25_map = {}
        if self.all_chunks:
            raw_scores = self.bm25.score(question)
            max_bm = max(raw_scores) if max(raw_scores) > 0 else 1
            indexed = sorted(enumerate(raw_scores), key=lambda x: -x[1])[:cfg.TOP_K * 3]
            for idx, s in indexed:
                key = self.all_chunks[idx][:100]
                if key not in bm25_map:
                    bm25_map[key] = {
                        "doc": Document(page_content=self.all_chunks[idx], metadata=self.all_metas[idx]),
                        "score": s / max_bm
                    }

        # ── Step 4: Fuse scores ──────────────────────────────────────
        all_keys = set(sem_map) | set(bm25_map)
        fused = []
        for key in all_keys:
            s_score = sem_map.get(key, {}).get("score", 0.0)
            b_score = bm25_map.get(key, {}).get("score", 0.0)
            combined = cfg.SEMANTIC_WEIGHT * s_score + cfg.BM25_WEIGHT * b_score
            doc = sem_map.get(key, bm25_map.get(key, {})).get("doc")
            if doc:
                fused.append((doc, combined, s_score, b_score))

        fused.sort(key=lambda x: -x[1])
        top_results = fused[:cfg.TOP_K]

        self.bts.update(
            num_retrieved=len(top_results),
            scores=[{"combined": round(f, 4), "semantic": round(s, 4), "bm25": round(b, 4)} for _, f, s, b in top_results],
            chunk_previews=[d.page_content[:120] for d, _, _, _ in top_results],
            sources=list({d.metadata.get("source", "?") for d, _, _, _ in top_results}),
            bm25_candidates=len(bm25_map), semantic_candidates=len(sem_map),
        )

        # ── Step 5: LLM generation ───────────────────────────────────
        context = "\n\n---\n\n".join(d.page_content for d, _, _, _ in top_results)
        prompt = ChatPromptTemplate.from_messages([
            ("system", "Answer the question using only the context below. Be clear and accurate. If the context doesn't contain the answer, say so honestly."),
            ("user", "Context:\n{context}\n\nQuestion: {question}")
        ])
        answer = (prompt | self.llm | StrOutputParser()).invoke({"context": context, "question": question})

        self.bts.update(context_chars=len(context), answer_chars=len(answer))

        # ── Step 6: Cache the result ─────────────────────────────────
        if cfg.CACHE_ENABLED:
            self.cache_store.add_documents([Document(
                page_content=question,
                metadata={"answer": answer, "sources": ", ".join(self.bts["sources"])}
            )])

        return {"answer": answer, "sources": ", ".join(self.bts["sources"]), "bts": self.bts}

    def get_stats(self) -> dict:
        try:
            return {
                "doc_chunks": self.vectorstore._collection.count(),
                "cached_queries": self.cache_store._collection.count()
            }
        except Exception:
            return {"doc_chunks": 0, "cached_queries": 0}
